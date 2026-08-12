# -*- coding: utf-8 -*-
"""Чтение вложений заявки Bitrix (воронка оплат 178) внутри фин.блока.
Тянет заявку по № → находит файловые поля → скачивает договор/КП/счёт →
читает через Anthropic API (PDF/сканы нативно) → извлекает условия для накопителя.

Куски перенесены из бота договоров (bitrix/attachments), адаптированы под воронку 178.
Env: BITRIX_WEBHOOK (тот же, что у sync_bitrix), ANTHROPIC_API_KEY (+ опц. CLAUDE_MODEL)."""
import base64, os, re, ssl, json, glob, hashlib, urllib.request

BASE = os.path.dirname(os.path.abspath(__file__))
_CTX = ssl.create_default_context(); _CTX.check_hostname = False; _CTX.verify_mode = ssl.CERT_NONE
CACHE = os.path.join(BASE, "data", "att_cache")


def _load_env():
    for name in (".env",):
        p = os.path.join(BASE, name)
        if os.path.exists(p):
            for line in open(p, encoding="utf-8"):
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, v = line.split("=", 1)
                    os.environ.setdefault(k.strip(), v.strip())


_load_env()
BITRIX = os.environ.get("BITRIX_WEBHOOK", "").rstrip("/")
BX_ENTITY = 178
BX_NUM_FIELD = "ufCrm4_1644310716"   # № заявки


def _bx(method, params):
    req = urllib.request.Request(f"{BITRIX}/{method}.json",
                                 data=json.dumps(params).encode("utf-8"),
                                 headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, context=_CTX, timeout=40) as r:
        d = json.load(r)
    if "error" in d:
        raise RuntimeError(f"{d['error']}: {d.get('error_description')}")
    return d["result"]


def item_by_num(num):
    r = _bx("crm.item.list", {"entityTypeId": BX_ENTITY, "filter": {BX_NUM_FIELD: str(num)}})
    items = r.get("items", [])
    return items[0] if items else None


def _file_entries(v):
    """Файловое UF-поле (dict или list of {urlMachine,url,downloadUrl,name}) → [(url,name)]."""
    entries = v if isinstance(v, list) else ([v] if v else [])
    out = []
    for e in entries:
        if isinstance(e, dict):
            u = e.get("urlMachine") or e.get("url") or e.get("downloadUrl")
            if u:
                out.append((u, e.get("name", "")))
    return out


def file_fields(item):
    """Все вложения карточки → [(field_code, url, name)] (авто-детект, без хардкода UF-кодов 178)."""
    out = []
    for k, v in (item or {}).items():
        for u, name in _file_entries(v):
            out.append((k, u, name))
    return out


_FIELD_LABELS = None


def field_labels():
    """Код поля → человекочитаемое название (Договор поставщика / АВР / Счёт / Тех требование …). Кэш."""
    global _FIELD_LABELS
    if _FIELD_LABELS is None:
        try:
            res = _bx("crm.item.fields", {"entityTypeId": BX_ENTITY})
            flds = res.get("fields", res) if isinstance(res, dict) else {}
            _FIELD_LABELS = {code: (f.get("title") or code) for code, f in (flds or {}).items()}
        except Exception:
            _FIELD_LABELS = {}
    return _FIELD_LABELS


def item_documents(item):
    """Прикреплённые документы заявки С ЯРЛЫКАМИ полей → [{field,label,name,url}].
    Ярлык говорит, ЧТО это (Договор/АВР/Счёт/Тех требование/Доверенность), — не гадаем."""
    labels = field_labels()
    out = []
    for field, url, name in file_fields(item):
        out.append({"field": field, "label": labels.get(field, field), "name": name, "url": url})
    return out


# какое поле-документ под какую задачу (по подстроке в НАЗВАНИИ поля) — для точечного ИИ-чтения
DOC_PURPOSE = {
    "article":   ("тех требован", "тех. требован", "техтребован"),   # вид работ → статья
    "contract":  ("договор",),                                        # условия (сумма/аванс/удержание)
    "avr":       ("авр", "накладн", "выполнен", "кс-2", "кс-3"),       # выполнено
    "invoice":   ("счёт", "счет"),                                     # сумма/№ счёта
}


def field_by_purpose(item, purpose):
    """Файлы заявки под задачу (article/contract/avr/invoice) → [пути скачивания] или []."""
    keys = DOC_PURPOSE.get(purpose, ())
    labels = field_labels()
    paths = []
    for field, url, name in file_fields(item):
        lbl = (labels.get(field, "") or "").lower()
        if any(k in lbl for k in keys):
            try:
                paths.append(download(url))
            except Exception:
                pass
    return paths


def read_purpose(item, purpose, instruction, schema):
    """Прочитать документ(ы) заявки ПОД ЗАДАЧУ (article/avr/contract/invoice) → dict, {} если файла нет."""
    paths = field_by_purpose(item, purpose)
    if not paths:
        return {}
    return read_docs(paths, instruction, schema)


# --- лёгкая схема: статья из Тех.требования (дёшево, для всех заявок) ---
_STR = {"type": "string"}
_NUM = {"type": "number"}
ARTICLE_JSON_SCHEMA = {
    "type": "object",
    "properties": {"article": _STR, "object": _STR, "ochered": _STR, "work_desc": _STR},
    "required": ["article", "object", "ochered", "work_desc"],
    "additionalProperties": False,
}
ARTICLE_INSTRUCTION = (
    "Перед тобой ТЕХНИЧЕСКОЕ ТРЕБОВАНИЕ / договор по заявке стройхолдинга (Казахстан). Определи: "
    "article = вид работ/поставки короткой фразой как в смете (Монолитные работы/Фундамент/Земляные/Кровля/"
    "Фасад/Внутренняя отделка/Окна и двери/Кладка стен/Инженерные сети/Лифт/Поставка материалов/Аренда техники/"
    "Проектные работы); object = ЖК (Атмосфера/Аура/Керуен/Аксай/…); ochered = очередь/блок; "
    "work_desc = что именно делают/поставляют, одной строкой.")

# --- АВР → выполнено (принятые работы, per заявка) ---
AVR_JSON_SCHEMA = {
    "type": "object",
    "properties": {"vypolneno_sum": _NUM, "act_no": _STR, "act_date": _STR,
                   "all_signed": {"type": "boolean"}, "notes": _STR},
    "required": ["vypolneno_sum", "act_no", "act_date", "all_signed", "notes"],
    "additionalProperties": False,
}
AVR_INSTRUCTION = (
    "Перед тобой АКТ ВЫПОЛНЕННЫХ РАБОТ (АВР / КС-2 / КС-3) или накладная. Извлеки: "
    "vypolneno_sum = сумма ПРИНЯТЫХ работ/поставки ИМЕННО по этому акту (число, без НДС-путаницы бери итог акта); "
    "act_no/act_date = № и дата акта; all_signed = подписан ли ВСЕМИ сторонами (заказчик+подрядчик, "
    "есть подписи и печати) true/false; notes — что принято, одной строкой.")

# --- счёт → сумма к оплате (для услуг/поставок сумма именно тут, а не в договоре) ---
INVOICE_JSON_SCHEMA = {
    "type": "object",
    "properties": {"total": _N, "account": _S, "account_date": _S, "nds": _S, "supplier": _S},
    "required": ["total", "account", "account_date", "nds", "supplier"],
    "additionalProperties": False,
}
INVOICE_INSTRUCTION = (
    "Перед тобой СЧЁТ НА ОПЛАТУ. Извлеки: total = ИТОГОВАЯ сумма к оплате по счёту (число, с НДС если "
    "итог с НДС); account = № счёта; account_date = дата счёта; nds = режим НДС (с НДС/без НДС); "
    "supplier = поставщик (кому платят).")


def download(url):
    """Скачать вложение → путь. Тип по magic-bytes. Кэш по md5 (повторно не качаем)."""
    os.makedirs(CACHE, exist_ok=True)
    key = hashlib.md5(url.encode()).hexdigest()[:16]
    hit = glob.glob(os.path.join(CACHE, key + ".*"))
    if hit:
        return hit[0]
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, context=_CTX, timeout=60) as r:
        data = r.read()
    head = data[:8]
    if head[:4] == b"%PDF":
        ext = "pdf"
    elif head[:2] == b"PK":
        ext = "docx" if (b"word/" in data[:6000] or b"word/document" in data) else "xlsx"
    elif head[:3] == b"\xff\xd8\xff":
        ext = "jpg"                        # Read определяет тип по расширению — даём настоящее
    elif head[:8] == b"\x89PNG\r\n\x1a\n":
        ext = "png"
    else:
        ext = "bin"
    p = os.path.join(CACHE, f"{key}.{ext}")
    open(p, "wb").write(data)
    return p


# ---------- чтение через Anthropic API ----------
def _docx_text(path):
    try:
        from docx import Document
        d = Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts += [c.text.strip() for c in row.cells if c.text.strip()]
        return "\n".join(parts)
    except Exception:
        return ""


def _xlsx_text(path):
    try:
        import openpyxl
        ws = openpyxl.load_workbook(path, data_only=True).worksheets[0]
        return "\n".join(" ".join(str(x) for x in row if x is not None)
                         for row in ws.iter_rows(values_only=True))
    except Exception:
        return ""


def _extract_json(text):
    """Выцепить JSON: снять ```-заборы, взять последнюю сбалансированную {…}."""
    if not text:
        return None
    t = re.sub(r"```(?:json)?", "", text)
    depth, start, best = 0, -1, None
    for i, ch in enumerate(t):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            if depth > 0:
                depth -= 1
                if depth == 0 and start >= 0:
                    best = t[start:i + 1]
    for cand in (best, t):
        if not cand:
            continue
        try:
            return json.loads(cand)
        except Exception:
            pass
    return None


def _build_content(paths, instruction):
    """Content-блоки для Messages API: image (сканы) / document (PDF) + текст (docx/xlsx)."""
    blocks, extra, has_media = [], "", False
    for p in paths:
        low = p.lower()
        try:
            if low.endswith(".pdf"):
                data = base64.standard_b64encode(open(p, "rb").read()).decode()
                blocks.append({"type": "document",
                               "source": {"type": "base64", "media_type": "application/pdf", "data": data}})
                has_media = True
            elif low.endswith((".png", ".jpg", ".jpeg")):
                media = "image/png" if low.endswith(".png") else "image/jpeg"
                data = base64.standard_b64encode(open(p, "rb").read()).decode()
                blocks.append({"type": "image", "source": {"type": "base64", "media_type": media, "data": data}})
                has_media = True
            elif low.endswith(".docx"):
                extra += f"\n[docx {os.path.basename(p)}]\n{_docx_text(p)[:12000]}"
            elif low.endswith(".xlsx"):
                extra += f"\n[xlsx {os.path.basename(p)}]\n{_xlsx_text(p)[:12000]}"
        except Exception:
            pass
    text = instruction + ("\n\nСодержимое приложенных документов:\n" + extra if extra else "")
    blocks.append({"type": "text", "text": text})
    return blocks, (has_media or bool(extra))


def read_docs(paths, instruction, schema_hint=None):
    """Прочитать документы через Anthropic API (PDF+сканы нативно) → dict условий.
    Структурированный вывод форсируется json_schema. Модель — CLAUDE_MODEL (по умолч. claude-opus-5)."""
    try:
        import anthropic
    except Exception:
        return {"error": "нет пакета anthropic (pip install anthropic)"}
    if not (os.environ.get("ANTHROPIC_API_KEY") or "").strip():
        return {"error": "ANTHROPIC_API_KEY не задан"}
    content, ok = _build_content(paths, instruction)
    if not ok:
        return {"error": "нет читаемых вложений"}
    model = os.environ.get("CLAUDE_MODEL", "claude-opus-5")
    try:
        client = anthropic.Anthropic()
        resp = client.messages.create(
            model=model, max_tokens=8000,
            output_config={"format": {"type": "json_schema", "schema": NAKOPITEL_JSON_SCHEMA},
                           "effort": "medium"},
            messages=[{"role": "user", "content": content}],
        )
    except Exception as e:
        return {"error": str(e)[:400]}
    if resp.stop_reason == "refusal":
        return {"error": "отказ модели (refusal)"}
    text = next((b.text for b in resp.content if b.type == "text"), "")
    try:
        return json.loads(text)          # output_config.format гарантирует валидный JSON
    except Exception:
        return _extract_json(text) or {"error": "не распознал JSON", "raw": (text or "")[:400]}


# строгая JSON-схема условий накопителя (structured outputs; additionalProperties=false обязателен)
_S = {"type": "string"}
_N = {"type": "number"}
NAKOPITEL_JSON_SCHEMA = {
    "type": "object",
    "properties": {
        "contract_no": _S, "contract_date": _S, "total": _N, "currency": _S, "nds": _S,
        "avans_pct": _N, "avans_sum": _N, "retention_pct": _N, "retention_sum": _N,
        "barter": {"type": "boolean"}, "barter_sum": _N,
        "ochered": _S, "object": _S, "account": _S, "bin": _S, "article": _S, "notes": _S,
    },
    "required": ["contract_no", "contract_date", "total", "currency", "nds", "avans_pct",
                 "avans_sum", "retention_pct", "retention_sum", "barter", "barter_sum",
                 "ochered", "object", "account", "bin", "article", "notes"],
    "additionalProperties": False,
}
NAKOPITEL_SCHEMA = NAKOPITEL_JSON_SCHEMA   # совместимость со старым импортом
NAKOPITEL_INSTRUCTION = (
    "Ты финансовый контролёр стройхолдинга (Казахстан). Перед тобой договор подряда/поставки и/или КП/счёт. "
    "Извлеки: № и дату договора; total = сумма ИМЕННО ЭТОГО договора/счёта (НЕ накопительный итог по объекту); "
    "валюту и режим НДС; аванс (% и сумму); гарантийное удержание (% и сумму); бартер (есть/нет и сумму); "
    "очередь/блок; объект (ЖК); account = № счёта/акта, если указан; "
    "bin = БИН/ИИН контрагента-подрядчика/поставщика (ровно 12 цифр из реквизитов сторон; "
    "НЕ ИИК/IBAN 'KZ…' и НЕ БИК). "
    "article = ВИД РАБОТ/поставки по договору короткой фразой как в строительной смете, напр.: "
    "'Монолитные работы', 'Фундамент', 'Земляные работы', 'Кровля', 'Фасад', 'Внутренняя отделка', "
    "'Окна и двери', 'Кладка стен', 'Инженерные сети', 'Лифтовое оборудование', 'Поставка арматуры/материалов', "
    "'Аренда техники', 'Проектные работы'. "
    "В notes — важные условия и оговорки одной строкой. Чего нет — ставь 0 или пустую строку.")
